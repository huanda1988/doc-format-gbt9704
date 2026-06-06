#!/usr/bin/env python3
"""Apply common GB/T 9704-2012 official-document layout to a DOCX.

This script handles the common no-seal document case: title, body, section
headings, final issuing authority, date, and alternating page-number footers.
For seals, attachments, imprint blocks, letter format, command format, or
minutes format, patch this script or build manually using the skill reference.
"""

from __future__ import annotations

import argparse
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


FANGSONG = "仿宋_GB2312"
HEITI = "黑体"
KAITI = "楷体_GB2312"
SONGTI = "宋体"
TITLE_FONT = "方正小标宋_GBK"

BODY_SIZE = Pt(16)  # 3号
TITLE_SIZE = Pt(22)  # 2号
PAGE_SIZE = Pt(14)  # 4号
LINE_PITCH = Pt(29)
CHAR_WIDTH = Pt(16)


def set_rfonts(element, font_name: str) -> None:
    rpr = element.get_or_add_rPr() if hasattr(element, "get_or_add_rPr") else element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        element.insert(0, rpr)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)


def set_style_font(style, font_name: str, size, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = size
    style.font.bold = bold
    set_rfonts(style.element, font_name)


def set_run_font(run, font_name: str, size, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    set_rfonts(run._element, font_name)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    return re.sub(r"\s+", " ", text).strip()


def paragraph_role(text: str) -> str:
    if re.match(r"^[一二三四五六七八九十]+、", text):
        return "h1"
    if re.match(r"^（[一二三四五六七八九十]+）", text):
        return "h2"
    if re.match(r"^\d+\.", text):
        return "h3"
    if re.match(r"^（\d+）", text):
        return "h4"
    return "body"


def restore_missing_first_section(texts: list[str]) -> list[str]:
    """Restore `一、` for a likely first-level heading when later h1s exist.

    Legacy .doc conversion sometimes drops an auto-numbered first section title
    while preserving later manually typed headings such as `二、...`.
    """
    if len(texts) < 4:
        return texts
    has_later_h1 = any(re.match(r"^[二三四五六七八九十]+、", t) for t in texts[1:])
    first_body_candidate_index = 2 if len(texts) > 2 else None
    if has_later_h1 and first_body_candidate_index is not None:
        candidate = texts[first_body_candidate_index]
        if (
            candidate
            and not paragraph_role(candidate).startswith("h")
            and len(candidate) <= 30
            and not candidate.endswith("。")
        ):
            texts[first_body_candidate_index] = "一、" + candidate
    return texts


def set_para_common(paragraph, first_indent: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.style = "Normal"
    fmt = paragraph.paragraph_format
    fmt.alignment = align
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_PITCH
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = CHAR_WIDTH * 2 if first_indent else Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)


def add_text_paragraph(
    doc: Document,
    text: str,
    font_name: str = FANGSONG,
    bold: bool = False,
    first_indent: bool = True,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
):
    p = doc.add_paragraph()
    set_para_common(p, first_indent=first_indent, align=align)
    run = p.add_run(text)
    set_run_font(run, font_name, BODY_SIZE, bold=bold)
    return p


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("— ")
    set_run_font(run, SONGTI, PAGE_SIZE)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)
    set_run_font(instr_run, SONGTI, PAGE_SIZE)

    fld_end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._r.append(fld_end)
    set_run_font(fld_end_run, SONGTI, PAGE_SIZE)

    run = paragraph.add_run(" —")
    set_run_font(run, SONGTI, PAGE_SIZE)


def enable_odd_even_footers(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def configure_footer(section) -> None:
    for footer, align, indent_side in (
        (section.footer, WD_ALIGN_PARAGRAPH.RIGHT, "right"),
        (section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT, "left"),
    ):
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(p)
        p.alignment = align
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(16)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if indent_side == "right":
            p.paragraph_format.right_indent = PAGE_SIZE
        else:
            p.paragraph_format.left_indent = PAGE_SIZE
        add_page_field(p)


def configure_section(section) -> None:
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.bottom_margin = Mm(35)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(28)
    configure_footer(section)

    sect_pr = section._sectPr
    for old in sect_pr.findall(qn("w:docGrid")):
        sect_pr.remove(old)
    doc_grid = OxmlElement("w:docGrid")
    doc_grid.set(qn("w:type"), "lines")
    doc_grid.set(qn("w:linePitch"), "580")
    sect_pr.append(doc_grid)


def extract_texts(source: Path) -> list[str]:
    doc = Document(source)
    return [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]


def build_document(source: Path, output: Path, keep_metadata: bool = False) -> Path:
    texts = restore_missing_first_section(extract_texts(source))
    if not texts:
        raise ValueError("Source DOCX has no readable paragraph text.")

    title = texts[0]
    signer = None
    date = None
    body = texts[1:]

    if len(texts) >= 3 and re.search(r"\d{4}年\d{1,2}月\d{1,2}日$", texts[-1]):
        signer = texts[-2]
        date = texts[-1]
        body = texts[1:-2]

    doc = Document()
    set_style_font(doc.styles["Normal"], FANGSONG, BODY_SIZE)
    normal_fmt = doc.styles["Normal"].paragraph_format
    normal_fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal_fmt.line_spacing = LINE_PITCH
    normal_fmt.space_before = Pt(0)
    normal_fmt.space_after = Pt(0)

    enable_odd_even_footers(doc)
    configure_section(doc.sections[0])

    title_p = doc.add_paragraph()
    set_para_common(title_p, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    title_p.paragraph_format.line_spacing = Pt(32)
    title_p.paragraph_format.space_after = LINE_PITCH
    title_run = title_p.add_run(title)
    set_run_font(title_run, TITLE_FONT, TITLE_SIZE)

    for text in body:
        role = paragraph_role(text)
        if role == "h1":
            add_text_paragraph(doc, text, font_name=HEITI)
        elif role == "h2":
            add_text_paragraph(doc, text, font_name=KAITI)
        else:
            add_text_paragraph(doc, text, font_name=FANGSONG)

    if signer and date:
        sign_p = add_text_paragraph(
            doc,
            signer,
            font_name=FANGSONG,
            first_indent=False,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        sign_p.paragraph_format.space_before = LINE_PITCH
        sign_p.paragraph_format.right_indent = CHAR_WIDTH * 4

        date_p = add_text_paragraph(
            doc,
            date,
            font_name=FANGSONG,
            first_indent=False,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        date_p.paragraph_format.right_indent = CHAR_WIDTH * 2

    if keep_metadata:
        src = Document(source)
        doc.core_properties.title = src.core_properties.title
        doc.core_properties.author = src.core_properties.author
        doc.core_properties.last_modified_by = src.core_properties.last_modified_by
    else:
        doc.core_properties.title = title
        doc.core_properties.author = ""
        doc.core_properties.last_modified_by = ""

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path, help="Input .docx file")
    parser.add_argument("--out", "-o", type=Path, required=True, help="Output .docx file")
    parser.add_argument("--keep-metadata", action="store_true", help="Preserve source core metadata")
    args = parser.parse_args()

    if args.input.suffix.lower() != ".docx":
        raise SystemExit("Input must be .docx. Convert .doc/.rtf/.txt to .docx first.")
    out = build_document(args.input, args.out, keep_metadata=args.keep_metadata)
    print(out)


if __name__ == "__main__":
    main()
